import asyncio
import os
import uuid

import chromadb
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.core.exceptions import BizError
from app.core.response import ResponseCode
from app.models import Document
from app.services.embedding_service import encode_texts

settings = get_settings()

SUPPORTED_EXTS = {"pdf", "docx", "md"}

_chroma_client: chromadb.api.ClientAPI | None = None
_chroma_lock = asyncio.Lock()


def get_chroma_client() -> chromadb.api.ClientAPI:
    global _chroma_client
    if _chroma_client is None:
        os.makedirs(settings.chroma_persist_dir, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=settings.chroma_persist_dir)
    return _chroma_client


def get_user_collection(user_id: int) -> chromadb.Collection:
    return get_chroma_client().get_or_create_collection(
        name=f"doc_user_{user_id}",
        metadata={"hnsw:space": "cosine"},
    )


def _parse_pdf_sync(file_path: str) -> str:
    import fitz

    doc = fitz.open(file_path)
    try:
        return "\n".join(page.get_text() for page in doc).strip()
    finally:
        doc.close()


def _parse_docx_sync(file_path: str) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs).strip()


def _parse_markdown_sync(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read().strip()


def _chunk_text_sync(text: str, chunk_size: int, overlap: int) -> list[str]:
    if not text:
        return []
    chunks = []
    start = 0
    step = max(chunk_size - overlap, 1)
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start += step
    return [c for c in chunks if c.strip()]


def chunk_text(text: str, chunk_size: int | None = None, overlap: int | None = None) -> list[str]:
    return _chunk_text_sync(
        text,
        chunk_size if chunk_size is not None else settings.chunk_size,
        overlap if overlap is not None else settings.chunk_overlap,
    )


async def parse_file(file_path: str, ext: str) -> str:
    parsers = {"pdf": _parse_pdf_sync, "docx": _parse_docx_sync, "md": _parse_markdown_sync}
    return await asyncio.to_thread(parsers[ext], file_path)


async def process_document(
    file_path: str,
    filename: str,
    ext: str,
    file_size: int,
    user_id: int,
    db: AsyncSession,
) -> Document:
    if ext not in SUPPORTED_EXTS:
        raise BizError(
            code=ResponseCode.UNSUPPORTED_FILE_TYPE,
            message=f"不支持的文件类型: {ext}",
            http_status=400,
        )

    text = await parse_file(file_path, ext)
    if not text:
        raise BizError(
            code=ResponseCode.DOC_PARSE_FAILED,
            message="无法解析文本内容（可能是扫描版 PDF，暂不支持 OCR）",
            http_status=400,
        )

    chunks = await asyncio.to_thread(_chunk_text_sync, text, settings.chunk_size, settings.chunk_overlap)
    if not chunks:
        raise BizError(code=ResponseCode.DOC_PARSE_FAILED, message="文档内容为空", http_status=400)

    embeddings = await encode_texts(chunks)

    document = Document(
        user_id=user_id,
        filename=filename,
        file_type=ext,
        chunk_count=len(chunks),
        file_size=file_size,
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    try:
        async with _chroma_lock:
            collection = get_user_collection(user_id)
            ids = [f"{document.id}_chunk_{i}" for i in range(len(chunks))]
            metadatas = [
                {
                    "user_id": user_id,
                    "document_id": document.id,
                    "filename": filename,
                    "chunk_index": i,
                }
                for i in range(len(chunks))
            ]
            collection.add(ids=ids, documents=chunks, embeddings=embeddings, metadatas=metadatas)
    except Exception:
        await db.delete(document)
        await db.commit()
        raise

    return document


async def delete_document(document: Document, db: AsyncSession) -> None:
    document_id = document.id
    user_id = document.user_id
    await db.delete(document)
    await db.commit()

    try:
        async with _chroma_lock:
            collection = get_user_collection(user_id)
            collection.delete(where={"document_id": document_id})
    except Exception:
        pass


def save_upload_file(content: bytes, ext: str) -> tuple[str, str]:
    os.makedirs("data/uploads", exist_ok=True)
    file_id = uuid.uuid4().hex[:12]
    file_path = os.path.abspath(os.path.join("data/uploads", f"{file_id}.{ext}"))
    with open(file_path, "wb") as f:
        f.write(content)
    return file_path, file_id
